# main.py
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pathlib import Path
from collections import Counter
import re, os, string, math, asyncio, time
from typing import List, Dict, Any

# Gemini Setup
import google.generativeai as genai
GEMINI_MODEL = "gemini-2.5-flash"
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    raise RuntimeError("GOOGLE_API_KEY belum diset.")
genai.configure(api_key=API_KEY)
_model = genai.GenerativeModel(GEMINI_MODEL)

# Ekspor
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Konfigurasi performa
MAX_CONCURRENCY = 10
MAX_RETRIES = 4
RETRY_BASE_DELAY = 0.8
GEMINI_TIMEOUT = 120
MAX_INPUT_CHARS = 50000

# ===================== ADAPTIVE FAST MODE CONFIG =====================
ADAPTIVE_PARAGRAPH_RULE = {
    "bab i": (5, 7),      # Pendahuluan
    "bab ii": (6, 10),    # Tinjauan Pustaka
    "bab iii": (5, 7),    # Metodologi
    "bab iv": (6, 12),    # Hasil & Pembahasan
    "bab v": (2, 4),      # Kesimpulan
}

def get_adaptive_paragraph_count(bab_title: str) -> int:
    bab = bab_title.lower().strip()
    for key, (min_p, max_p) in ADAPTIVE_PARAGRAPH_RULE.items():
        if key in bab:
            return max_p  # pakai batas maksimum biar isi lebih lengkap
    return 5  # default aman

def clean_bab_title(title: str) -> str:
    return title.replace("–", "-").replace("…", "").strip()

# Ekstraksi PDF
def read_pdf_text(path: str) -> str:
    text = ""
    try:
        import fitz
        doc = fitz.open(path)
        text = "\n".join([p.get_text() for p in doc])
        text = clean_text(text)
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from pdfminer.high_level import extract_text
        text = clean_text(extract_text(path))
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        raw = "\n".join([p.extract_text() or "" for p in reader.pages])
        text = clean_text(raw)
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from pdf2image import convert_from_path
        import pytesseract
        pages = convert_from_path(path, dpi=300)
        ocr_text = ""
        for pg in pages:
            ocr_text += pytesseract.image_to_string(pg, lang="eng+ind") + "\n"
        return clean_text(ocr_text)
    except:
        return clean_text(text)

def _enough_text(text, min_chars=200):
    return len(text.strip()) >= min_chars

# Cleaning Text
BLACKBOX = ["■","□","▯","█","�"]

def clean_text(text):
    if not text:
        return ""
    for b in BLACKBOX:
        text = text.replace(b, "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Gambar\s*\d+(\.\d+)*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"Tabel\s*\d+(\.\d+)*", "", text, flags=re.IGNORECASE)

    return text.strip()

def clean_reference_noise(text):
    # Hapus URL
    text = re.sub(r"http\S+|www\S+", "", text)

    # Hapus pola (Nama, 2019)
    text = re.sub(r"\([A-Za-z][^()]{0,40}\d{4}\)", "", text)

    # Hapus Nama, 2019
    text = re.sub(r"[A-Za-z]+,\s*\d{4}", "", text)

    # Hapus daftar pustaka multiline
    text = re.sub(r"([A-Za-z]+\s*,){2,}.*", "", text)

    # Hapus institusi yang berulang
    text = re.sub(r"(Universitas|Fakultas|Program Studi|Jurusan|Departemen).*", "", text)

    # Rapikan spasi
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

# FIX SPLIT BAB 
def split_by_bab(text: str):
    # Buang DAFTAR ISI dan elemen awal
    text = re.sub(r"DAFTAR\s+ISI.*?(?=BAB\s+I\b|BAB\s+1\b)", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"DAFTAR\s+GAMBAR.*?(?=BAB\s+I\b|BAB\s+1\b)", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"DAFTAR\s+TABEL.*?(?=BAB\s+I\b|BAB\s+1\b)", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"DAFTAR PUSTAKA.*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"LAMPIRAN.*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"Tabel\s*\d+.*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"Gambar\s*\d+.*", "", text, flags=re.IGNORECASE)

    # Hapus baris titik-titik daftar isi
    text = re.sub(r"^.*\.{5,}.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\([A-Za-z].{0,50}\d{4}\)", "", text)  # (Nama, tahun)
    text = re.sub(r"[A-Za-z]+,\s*\d{4}", "", text) 
    text = re.sub(r"http\S+|www\S+", "", text)

    # Hapus nomor halaman Romawi
    text = re.sub(r"(?m)^\s*[ivxlcdm]+\s*$", "", text, flags=re.IGNORECASE)

    # Cari mulai dari BAB I atau BAB 1
    m = re.search(r"(BAB\s+(?:I|1)\b.*)", text, flags=re.IGNORECASE | re.DOTALL)
    if m:
        text = m.group(1)

    # Pecah BAB versi lama (romawi)
    parts = re.split(r"(?=BAB\s+[IVXLCDM]+\b)", text, flags=re.IGNORECASE)

    # Fallback: kalau tidak ada romawi, pakai angka
    if len(parts) <= 1:
        parts = re.split(r"(?=BAB\s+\d+\b)", text, flags=re.IGNORECASE)

    hasil = []
    for p in parts:
        p = p.strip()
        if not p:
            continue

        # terima BAB I/II/III dan BAB 1/2/3
        if not re.match(r"^BAB\s+(?:[IVXLCDM]+|\d+)\b", p, flags=re.IGNORECASE):
            continue

        lines = p.split("\n", 1)
        if len(lines) < 2:
            continue

        judul = lines[0].strip()
        isi = lines[1].strip()

        isi = re.sub(r"^.{5,100}\.{3,}\d+$", "", isi, flags=re.MULTILINE)

        hasil.append({"judul": judul, "isi": isi})

    return hasil

# Utility tokenisasi & ekstraktif
import string
from collections import Counter
import math

STOPWORDS = set("yang dan di ke dari untuk pada adalah dengan dalam ini itu serta juga tidak dapat atau oleh bagi agar sudah akan para sebagai tersebut karena maka sehingga terhadap serta olehnya".split())
PUNCT = str.maketrans("", "", string.punctuation)

def tokenize(text: str):
    return [w for w in text.lower().translate(PUNCT).split() if w not in STOPWORDS and len(w) > 2]

def split_sentences(text: str):
    sents = re.split(r"(?<=[\.\?\!])\s+(?=[A-Za-z0-9])", text.strip())
    return [s.strip() for s in sents if s.strip()]

def summarize_text_extractive(text: str, max_sent: int = 8) -> str:
    sents = split_sentences(text)
    if not sents:
        return ""
    sent_tokens = [tokenize(s) for s in sents]
    df = Counter()
    for t in sent_tokens:
        df.update(set(t))
    N = len(sents)
    scores = []
    for i, toks in enumerate(sent_tokens):
        score = sum(
            (cnt / (1 + len(toks))) * (math.log((N + 1) / (1 + df[w])) + 1)
            for w, cnt in Counter(toks).items()
        )
        if i < max(3, int(N * 0.1)):
            score *= 1.15
        scores.append(score)
    top_idx = sorted(range(N), key=lambda i: scores[i], reverse=True)[:max_sent]
    return " ".join([sents[i] for i in sorted(top_idx)])

# Prompt Gemini
# ===================== ADAPTIVE GEMINI PROMPT =====================
SUM_PROMPT_TEMPLATE = """
Tugas kamu adalah merangkum isi {bab_title} dari dokumen ilmiah (skripsi) secara akademik.
Ringkasan harus disusun runtut sesuai konteks isi bab dan tidak hanya mengambil bagian awal saja.
Gunakan bahasa ilmiah yang natural seperti skripsi kampus dan tidak terdeteksi AI.

WAJIB:
- {max_paragraphs} paragraf sesuai aturan BAB
- Bahasa ilmiah formal
- Tidak menambah teori baru
- Tidak mengubah makna isi
- Tidak copy paste kalimat

HAPUS:
- Daftar pustaka, referensi, (Nama, Tahun)
- URL/link
- Nomor tabel/gambar/lampiran
- Isi yang tidak relevan

TEKS SUMBER (TERPILIH):
\"\"\"{content}\"\"\"

HASIL RINGKASAN:
"""

# ===================== ADAPTIVE PARAGRAPH CONTROL PER BAB =====================
def estimate_paragraph_count(bab_title: str, content: str) -> int:
    length = len(content)

    if "bab i" in bab_title.lower():
        return 5 if length < 8000 else 7
    elif "bab ii" in bab_title.lower():
        return 6 if length < 12000 else 10
    elif "bab iii" in bab_title.lower():
        return 5 if length < 9000 else 7
    elif "bab iv" in bab_title.lower():
        return 6 if length < 15000 else 12
    elif "bab v" in bab_title.lower():
        return 2 if length < 5000 else 4
    else:
        return 4  # fallback aman

def _gemini_call_sync(prompt: str) -> str:
    resp = _model.generate_content(
        prompt,
        request_options={"timeout": GEMINI_TIMEOUT},
        generation_config={"response_mime_type": "text/plain"},
    )
    return (getattr(resp, "text", "") or "").strip()

import asyncio, time

async def gemini_summarize_async(content: str, semaphore: asyncio.Semaphore, bab_title: str):
    max_paragraphs = estimate_paragraph_count(bab_title, content)

    prompt = SUM_PROMPT_TEMPLATE.format(
        bab_title=bab_title,
        content=content,
        max_paragraphs=max_paragraphs 
    )

    attempt = 0
    while True:
        attempt += 1
        try:
            async with semaphore:
                result = await asyncio.to_thread(_gemini_call_sync, prompt)

            if result and len(result.strip()) > 120:
                return result.strip()  # berhasil
            else:
                raise ValueError("Gemini output terlalu pendek, retry...")
        except Exception:
            if attempt >= MAX_RETRIES:
                return summarize_text_extractive(content, max_sent=8)  # fallback cepat
            await asyncio.sleep(RETRY_BASE_DELAY * attempt)

# Pre-Compress per BAB
def compress_for_prompt(text: str, max_chars: int = MAX_INPUT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    base_k = 10 + min(4, len(text) // 20000)  # 10–14 kalimat
    extract = summarize_text_extractive(text, max_sent=base_k)
    if len(extract) > max_chars:
        extract = extract[:max_chars]
    return extract

def hard_clean_output(text: str) -> str:
    # Buang header sampah yang sering kebawa
    text = re.sub(r"(?i)^(DAFTAR\s+ISI|BAB\s+[IVXLCDM]+|HALAMAN\s+JUDUL|KATA\s+PENGANTAR|LEMBAR\s+PENGESAHAN).*", "", text, flags=re.MULTILINE)
    # Buang titik titik daftar isi
    text = re.sub(r"\.{5,}", "", text)
    # Hapus nomor halaman sisa
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    # Hapus tanda kurung referensi sisa [1], (2), dsb
    text = re.sub(r"\[\d+\]|\(\d+\)", "", text)
    # Rapikan spasi
    text = re.sub(r"[ ]{2,}", " ", text)
    # Rapikan newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Hilangkan spasi sebelum tanda baca
    text = re.sub(r" \,", ",", text)
    text = re.sub(r" \.", ".", text)
    text = re.sub(r" \;", ";", text)
    text = re.sub(r" \:", ":", text)
    return text.strip()

# ===================== AUTO FILTER SPAM & DAFTAR ISI =====================
def auto_filter_noise(text: str) -> str:
    banned_patterns = [
        r"daftar isi", r"lembar pengesahan", r"abstrak", r"kata pengantar",
        r"lembar persetujuan", r"pernyataan keaslian", r"universitas", r"fakultas",
        r"program studi", r"jurusan", r"pembimbing", r"nim", r"nip", r"dosen",
        r"lampiran", r"daftar pustaka", r"universitas .*?\n", r"bab i pendahuluan"


    ]
    for pattern in banned_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    return text.strip()

# ===================== AUTO MERGE SHORT SENTENCES =====================
def merge_short_sentences(text: str, min_len: int = 40) -> str:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    merged = []
    buffer = ""

    for s in sentences:
        if len(s) < min_len:
            buffer += " " + s
        else:
            if buffer:
                merged.append((buffer + " " + s).strip())
                buffer = ""
            else:
                merged.append(s.strip())
    if buffer:
        merged.append(buffer.strip())

    return " ".join(merged)

# ===================== HUMANIZE =====================
async def humanize_final(text: str, semaphore):
    prompt = f"""
    Ubah teks berikut agar lebih alami seperti tulisan manusia namun tetap formal akademik.
    Jangan ubah makna ilmiah, cukup variasikan struktur kalimat agar tidak terdeteksi AI
    dan tidak mirip sumber asli.

    Aturan:
    - Jangan gunakan gaya AI seperti: "Selain itu", "Secara keseluruhan", "Dengan demikian"
    - Hindari pengulangan frasa
    - Perbaiki alur antar kalimat agar mengalir
    - Gaya bahasa seperti skripsi kampus Indonesia

    Teks:
    \"\"\"{text}\"\"\"

    Versi final natural dan aman:
    """
    try:
        async with semaphore:
            return await asyncio.to_thread(_gemini_call_sync, prompt)
    except:
        return text

# Ringkas Per BAB (paralel)
async def summarize_sections_parallel(sections: List[Dict[str, str]]) -> List[Dict[str, Any]]:
    semaphore = asyncio.Semaphore(MAX_CONCURRENCY)

    async def _process(sec):
        isi = (sec.get("isi") or "").strip()
        paragraf = [p.strip() for p in isi.split("\n") if p.strip()]
        paragraf_bersih = [p for p in paragraf if len(p) > 40]

        if len(paragraf_bersih) >= 2:
            isi_bersih = "\n".join(paragraf_bersih[:70])
        else:
            return {"judul": sec.get("judul", ""), "ringkasan_bab": ""}

        isi_bersih = clean_reference_noise(isi_bersih)
        compressed = compress_for_prompt(isi_bersih, MAX_INPUT_CHARS)

        summary = await gemini_summarize_async(compressed, semaphore, sec.get("judul", ""))

        summary = hard_clean_output(summary)  
        summary = auto_filter_noise(summary)
        summary = merge_short_sentences(summary)

        summary = await humanize_final(summary, semaphore)

        return {"judul": sec.get("judul", ""), "ringkasan_bab": summary}

    tasks = [asyncio.create_task(_process(sec)) for sec in sections]
    return await asyncio.gather(*tasks)

# Driver: ringkas PDF per BAB
async def summarize_pdf_per_bab(path: str):
    raw = read_pdf_text(path)
    # Safety: kalau extractor gagal atau hasil kosong, balikin struktur kosong rapi
    if not raw.strip():
        return {"file": os.path.basename(path), "sections": []}

    sections = split_by_bab(raw)

    # Safety: kalau split gagal, pakai 1 “BAB I” wildcard supaya backend tetap jalan
    if not sections:
        sections = [{"judul": "BAB I", "isi": raw}]

    results = await summarize_sections_parallel(sections)
    return {"file": os.path.basename(path), "sections": results}

# Export DOCX & PDF
def export_all(data, out_docx, out_pdf):
    # DOCX
    doc = Document()
    doc.add_heading("Ringkasan Per Bab (Gemini 2.5 Flash)", 0)
    doc.add_paragraph(f"File: {data['file']}")
    for sec in data["sections"]:
        doc.add_heading(sec["judul"], level=1)
        doc.add_paragraph(sec["ringkasan_bab"] or "")
    doc.save(out_docx)

    # PDF
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(out_pdf, pagesize=A4)
    elements = []
    elements.append(Paragraph("Ringkasan Per Bab", styles['Title']))
    elements.append(Paragraph(f"File: {data['file']}", styles['Normal']))
    elements.append(Spacer(1, 12))
    for sec in data["sections"]:
        elements.append(Paragraph(sec["judul"], styles['Heading2']))
        elements.append(Paragraph(sec["ringkasan_bab"] or "", styles['Normal']))
        elements.append(Spacer(1, 12))
    pdf.build(elements)

# FastAPI
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

app = FastAPI(
    title="DocuSum AI (Gemini) — Balanced Fast",
    description="Ringkasan Per Bab rapi (PDF, paralel, 4-5 paragraf, bahasa mengikuti dokumen.",
    version="8.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # batasi di production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
BASE_URL = "https://docusum.onrender.com"

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Hanya file PDF diperbolehkan.")
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())

    try:
        hasil = await summarize_pdf_per_bab(str(file_path))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal meringkas: {e}")

    docx_path = str(file_path.with_suffix(".docx"))
    pdf_path = str(file_path.with_suffix(".summary.pdf"))
    try:
        export_all(hasil, docx_path, pdf_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal ekspor: {e}")

    return {
        "success": True,
        "download_docx": f"{BASE_URL}/api/download/{Path(docx_path).name}",
        "download_pdf": f"{BASE_URL}/api/download/{Path(pdf_path).name}",
        "data": hasil
    }

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File tidak ditemukan")
    return FileResponse(file_path, filename=filename)

# ===================== KOMENTAR GLOBAL (TERSIMPAN PERMANEN) =====================
import json
from datetime import datetime


COMMENTS_FILE = Path("comments.json")

def load_comments() -> list:
    if COMMENTS_FILE.exists():
        try:
            with open(COMMENTS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except json.JSONDecodeError:
            return []
    return []

def save_comments(comments: list):
    with open(COMMENTS_FILE, "w", encoding="utf-8") as f:
        json.dump(comments, f, ensure_ascii=False, indent=2)

@app.get("/api/comments")
async def get_comments():
    """Ambil semua komentar yang tersimpan"""
    return load_comments()

@app.post("/api/comments")
async def post_comment(comment: Dict[str, str]):
    """Tambah komentar baru dan simpan permanen"""
    name = (comment.get("name") or "").strip()
    text = (comment.get("text") or "").strip()


    if not text:
        raise HTTPException(status_code=400, detail="Komentar tidak boleh kosong.")

    new_comment = {
        "name": name or "Anonim",
        "text": text,
        "time": datetime.utcnow().isoformat() + "Z"
    }

    comments = load_comments()
    comments.append(new_comment)
    save_comments(comments)

    return {"success": True, "comment": new_comment}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
